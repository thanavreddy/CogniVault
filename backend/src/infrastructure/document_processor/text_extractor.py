"""Extract plain text from PDF, DOCX, TXT, and Markdown files."""
import logging
import os
from pathlib import Path

from src.domain.entities.document import DocumentType

logger = logging.getLogger(__name__)


class TextExtractor:
    """Extracts clean text from various document formats."""

    def extract(self, file_path: str, document_type: DocumentType) -> str:
        """Extract text from a file based on its type."""
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")

        extractors = {
            DocumentType.PDF: self.extract_pdf,
            DocumentType.DOCX: self.extract_docx,
            DocumentType.TXT: self.extract_txt,
            DocumentType.MARKDOWN: self.extract_markdown,
        }

        extractor = extractors.get(document_type)
        if not extractor:
            raise ValueError(f"Unsupported document type: {document_type}")

        text = extractor(file_path)
        logger.info(
            "Extracted %d characters from %s (type=%s)",
            len(text), file_path, document_type.value,
        )
        return self._clean_text(text)

    def extract_pdf(self, file_path: str) -> str:
        """Extract text from a PDF file using pypdf."""
        try:
            from pypdf import PdfReader
            reader = PdfReader(file_path)
            pages = []
            for page_num, page in enumerate(reader.pages, 1):
                page_text = page.extract_text()
                if page_text:
                    pages.append(f"[Page {page_num}]\n{page_text}")
            return "\n\n".join(pages)
        except ImportError:
            logger.warning("pypdf not available, falling back to unstructured")
            return self._extract_with_unstructured(file_path)

    def extract_docx(self, file_path: str) -> str:
        """Extract text from a DOCX file."""
        try:
            from docx import Document
            doc = Document(file_path)
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            
            # Also extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text for cell in row.cells if cell.text.strip())
                    if row_text:
                        paragraphs.append(row_text)
            
            return "\n\n".join(paragraphs)
        except ImportError:
            return self._extract_with_unstructured(file_path)

    def extract_txt(self, file_path: str) -> str:
        """Read plain text file."""
        with open(file_path, "r", encoding="utf-8", errors="replace") as f:
            return f.read()

    def extract_markdown(self, file_path: str) -> str:
        """Read markdown file (strip markdown syntax for embedding)."""
        with open(file_path, "r", encoding="utf-8", errors="replace") as f:
            content = f.read()
        # Basic markdown stripping for cleaner embeddings
        import re
        content = re.sub(r"^#{1,6}\s+", "", content, flags=re.MULTILINE)  # Headers
        content = re.sub(r"\*\*(.+?)\*\*", r"\1", content)  # Bold
        content = re.sub(r"\*(.+?)\*", r"\1", content)  # Italic
        content = re.sub(r"`{1,3}[^`]*`{1,3}", "", content)  # Code
        content = re.sub(r"!\[.*?\]\(.*?\)", "", content)  # Images
        content = re.sub(r"\[(.+?)\]\(.*?\)", r"\1", content)  # Links
        return content

    def _extract_with_unstructured(self, file_path: str) -> str:
        """Fallback extraction using unstructured library."""
        try:
            from unstructured.partition.auto import partition
            elements = partition(filename=file_path)
            return "\n\n".join([str(el) for el in elements])
        except ImportError:
            logger.error("unstructured library not available")
            raise

    @staticmethod
    def _clean_text(text: str) -> str:
        """Remove excessive whitespace and normalize line endings."""
        import re
        text = text.replace("\r\n", "\n").replace("\r", "\n")
        text = re.sub(r"\n{3,}", "\n\n", text)  # Max 2 consecutive newlines
        text = re.sub(r"[ \t]+", " ", text)      # Normalize spaces
        return text.strip()
